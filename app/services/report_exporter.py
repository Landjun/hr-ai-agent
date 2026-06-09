"""把报告 Markdown 导出为 HTML 与 PDF。

- HTML：用 markdown 库渲染 + 内置清爽样式，浏览器打开即可阅读 / Ctrl+P 存 PDF。
- PDF ：用 reportlab + 内置 CJK 字体 STSong-Light，无需外部字体、跨平台，中文正常。
"""
from __future__ import annotations

import html as _html
import io
import re
from typing import List

# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------
_HTML_CSS = """
:root { color-scheme: light; }
body { font-family: "Microsoft YaHei","PingFang SC","Hiragino Sans GB",
       "Source Han Sans SC","Noto Sans CJK SC",sans-serif;
       max-width: 820px; margin: 32px auto; padding: 0 20px; color: #1f2328;
       line-height: 1.7; }
h1 { font-size: 26px; border-bottom: 2px solid #2563eb; padding-bottom: 8px; }
h2 { font-size: 20px; margin-top: 28px; color: #1d4ed8; }
h3 { font-size: 16px; margin-top: 18px; }
table { border-collapse: collapse; width: 100%; margin: 12px 0; }
th, td { border: 1px solid #d0d7de; padding: 6px 10px; text-align: left;
         font-size: 14px; vertical-align: top; }
th { background: #f3f6fb; }
ul { padding-left: 22px; }
blockquote { color: #57606a; border-left: 4px solid #d0d7de; margin: 12px 0;
             padding: 4px 14px; background: #f8fafc; }
hr { border: none; border-top: 1px solid #e5e7eb; margin: 24px 0; }
code { background:#f3f4f6; padding:1px 5px; border-radius:4px; }
@media print { body { margin: 0; } }
"""


def markdown_to_html(md: str, title: str = "报告") -> str:
    """Markdown → 完整可打印 HTML 文档。"""
    try:
        import markdown as _md
        body = _md.markdown(md, extensions=["tables", "sane_lists", "nl2br"])
    except Exception:
        body = "<pre>" + _html.escape(md) + "</pre>"
    return (f"<!doctype html><html lang=\"zh\"><head><meta charset=\"utf-8\">"
            f"<meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">"
            f"<title>{_html.escape(title)}</title><style>{_HTML_CSS}</style></head>"
            f"<body>{body}</body></html>")


# ---------------------------------------------------------------------------
# PDF（reportlab + STSong-Light）
# ---------------------------------------------------------------------------
_FONT = "STSong-Light"
_font_ready = False


def _ensure_font() -> None:
    global _font_ready
    if _font_ready:
        return
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    pdfmetrics.registerFont(UnicodeCIDFont(_FONT))
    _font_ready = True


def _inline(text: str) -> str:
    """转义 + 处理 **加粗** / *斜体*，供 reportlab Paragraph 使用。"""
    text = _html.escape(text, quote=False)
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*([^*]+)\*", r"<i>\1</i>", text)
    return text


def markdown_to_pdf_bytes(md: str, title: str = "报告") -> bytes:
    """把（本系统生成的）报告 Markdown 渲染成 PDF 字节。"""
    _ensure_font()
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer,
                                    Table, TableStyle)

    body = ParagraphStyle("body", fontName=_FONT, fontSize=10.5, leading=16)
    h1 = ParagraphStyle("h1", fontName=_FONT, fontSize=19, leading=24,
                        textColor=colors.HexColor("#1d4ed8"), spaceAfter=8)
    h2 = ParagraphStyle("h2", fontName=_FONT, fontSize=14, leading=20,
                        textColor=colors.HexColor("#1d4ed8"), spaceBefore=10, spaceAfter=4)
    h3 = ParagraphStyle("h3", fontName=_FONT, fontSize=12, leading=17,
                        spaceBefore=6, spaceAfter=2)
    bullet = ParagraphStyle("bullet", parent=body, leftIndent=14, bulletIndent=4)
    quote = ParagraphStyle("quote", parent=body, leftIndent=10,
                          textColor=colors.HexColor("#57606a"))
    cell = ParagraphStyle("cell", fontName=_FONT, fontSize=9, leading=13)

    flow: List = []
    lines = md.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        # 表格：连续以 | 开头的行
        if stripped.startswith("|"):
            tbl: List[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl.append(lines[i].strip())
                i += 1
            flow.append(_make_table(tbl, cell))
            flow.append(Spacer(1, 4))
            continue

        if not stripped:
            flow.append(Spacer(1, 5))
        elif stripped.startswith("# "):
            flow.append(Paragraph(_inline(stripped[2:]), h1))
        elif stripped.startswith("## "):
            flow.append(Paragraph(_inline(stripped[3:]), h2))
        elif stripped.startswith("### "):
            flow.append(Paragraph(_inline(stripped[4:]), h3))
        elif stripped.startswith(("- ", "* ")):
            flow.append(Paragraph("• " + _inline(stripped[2:]), bullet))
        elif stripped.startswith(">"):
            flow.append(Paragraph(_inline(stripped.lstrip("> ").strip()), quote))
        elif set(stripped) <= {"-", "*", "_"} and len(stripped) >= 3:
            flow.append(Spacer(1, 6))  # 分隔线
        else:
            flow.append(Paragraph(_inline(stripped), body))
        i += 1

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=18 * mm, bottomMargin=18 * mm,
                            leftMargin=16 * mm, rightMargin=16 * mm, title=title)
    doc.build(flow)
    return buf.getvalue()


def _make_table(rows: List[str], cell_style):
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, Table, TableStyle

    parsed = []
    for r in rows:
        cells = [c.strip() for c in r.strip().strip("|").split("|")]
        # 跳过 |---|---| 分隔行
        if cells and all(set(c) <= {"-", ":", " "} and c for c in cells):
            continue
        parsed.append([Paragraph(_inline(c), cell_style) for c in cells])
    if not parsed:
        return Paragraph("", cell_style)
    tbl = Table(parsed, hAlign="LEFT", repeatRows=1)
    tbl.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d0d7de")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f3f6fb")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
    ]))
    return tbl


def title_from_markdown(md: str, default: str = "报告") -> str:
    for line in md.splitlines():
        if line.strip().startswith("# "):
            return line.strip()[2:].strip()
    return default


# ---------------------------------------------------------------------------
# Word (.docx)
# ---------------------------------------------------------------------------
def _add_runs(paragraph, text: str) -> None:
    """把 **加粗** 文本拆成 run 写入段落（python-docx 原生支持中文）。"""
    for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = (i % 2 == 1)  # 奇数段是 ** 之间的内容


def markdown_to_docx_bytes(md: str, title: str = "报告") -> bytes:
    """把（本系统生成的）报告 Markdown 渲染成 Word 文档字节。"""
    import docx
    from docx.shared import Pt

    doc = docx.Document()
    doc.core_properties.title = title
    lines = md.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        s = lines[i].strip()

        if s.startswith("|"):  # 表格
            tbl_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not (cells and all(set(c) <= {"-", ":", " "} and c for c in cells)):
                    tbl_rows.append(cells)
                i += 1
            if tbl_rows:
                cols = max(len(r) for r in tbl_rows)
                table = doc.add_table(rows=0, cols=cols)
                table.style = "Light Grid Accent 1"
                for ri, row in enumerate(tbl_rows):
                    cells = table.add_row().cells
                    for ci in range(cols):
                        val = row[ci] if ci < len(row) else ""
                        cells[ci].text = ""
                        _add_runs(cells[ci].paragraphs[0], val)
                        if ri == 0:
                            for run in cells[ci].paragraphs[0].runs:
                                run.bold = True
            continue

        if not s:
            pass
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=0)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=1)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=2)
        elif s.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, s[2:])
        elif s.startswith(">"):
            p = doc.add_paragraph()
            run = p.add_run(s.lstrip("> ").strip())
            run.italic = True
        elif set(s) <= {"-", "*", "_"} and len(s) >= 3:
            pass  # 分隔线忽略
        else:
            _add_runs(doc.add_paragraph(), s)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
