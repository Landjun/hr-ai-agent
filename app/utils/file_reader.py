"""多格式简历读取：PDF / DOCX / TXT / MD / 直接粘贴文本。

读取失败时不抛异常中断，而是返回带提示的占位文本，保证流水线不崩。
"""
from __future__ import annotations

from pathlib import Path
from typing import Union

from app.utils.text_cleaner import clean_text


def read_pdf(path: Union[str, Path]) -> str:
    text_parts: list[str] = []
    # 优先 pdfplumber（排版更准），失败回退 pypdf
    try:
        import pdfplumber

        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text_parts.append(page.extract_text() or "")
        if any(p.strip() for p in text_parts):
            return "\n".join(text_parts)
    except Exception:
        pass

    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        for page in reader.pages:
            text_parts.append(page.extract_text() or "")
        return "\n".join(text_parts)
    except Exception as exc:
        return f"（PDF 解析失败：{exc}）"


def read_docx(path: Union[str, Path]) -> str:
    try:
        import docx

        doc = docx.Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" ".join(c.text for c in row.cells))
        return "\n".join(parts)
    except Exception as exc:
        return f"（DOCX 解析失败：{exc}）"


def read_text_file(path: Union[str, Path]) -> str:
    for enc in ("utf-8", "gbk", "utf-16"):
        try:
            return Path(path).read_text(encoding=enc)
        except Exception:
            continue
    return Path(path).read_text(encoding="utf-8", errors="ignore")


def read_file(path: Union[str, Path]) -> str:
    """按扩展名分发读取，返回清洗后的纯文本。"""
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        raw = read_pdf(p)
    elif suffix in (".docx", ".doc"):
        raw = read_docx(p)
    elif suffix in (".txt", ".md", ".markdown"):
        raw = read_text_file(p)
    else:
        raw = read_text_file(p)  # 兜底当文本读
    return clean_text(raw)


def read_bytes(data: bytes, file_name: str) -> str:
    """读取上传的文件字节流（Streamlit / FastAPI 上传场景）。"""
    import tempfile

    suffix = Path(file_name).suffix or ".txt"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        return read_file(tmp_path)
    finally:
        try:
            Path(tmp_path).unlink()
        except Exception:
            pass
